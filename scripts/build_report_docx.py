# -*- coding: utf-8 -*-
from __future__ import annotations

import math
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report_assets"
OUT = DOCS / "bao-cao-do-an-tot-nghiep-p2p-chat.docx"

NAVY = "#12365f"
TEAL = "#0f766e"
MINT = "#dff7ef"
BLUE = "#2563eb"
ORANGE = "#f59e0b"
RED = "#dc2626"
GREEN = "#16a34a"
GRAY = "#64748b"
LIGHT = "#f8fafc"
BORDER = "#cbd5e1"


def font_path(*names: str) -> str | None:
    for name in names:
        p = Path("C:/Windows/Fonts") / name
        if p.exists():
            return str(p)
    return None


FONT_REG = font_path("arial.ttf", "segoeui.ttf")
FONT_BOLD = font_path("arialbd.ttf", "segoeuib.ttf")
FONT_MONO = font_path("consola.ttf", "cour.ttf")


def pil_font(size: int, bold: bool = False, mono: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_MONO if mono else (FONT_BOLD if bold else FONT_REG)
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        words = raw.split()
        current = ""
        for word in words:
            candidate = word if not current else f"{current} {word}"
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        lines.append(current)
    return [line for line in lines if line]


def centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = "#0f172a",
    max_width: int | None = None,
    line_gap: int = 8,
) -> None:
    x1, y1, x2, y2 = box
    width = max_width or (x2 - x1 - 28)
    lines = wrap(draw, text, font, width)
    metrics = [draw.textbbox((0, 0), line, font=font) for line in lines]
    heights = [m[3] - m[1] for m in metrics]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, bbox in zip(lines, metrics):
        w = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font, fill=fill)
        y += (bbox[3] - bbox[1]) + line_gap


def box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str = "",
    fill: str = "white",
    outline: str = BORDER,
    title_color: str = NAVY,
    radius: int = 26,
) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    title_font = pil_font(30, bold=True)
    body_font = pil_font(23)
    draw.text((x1 + 24, y1 + 20), title, font=title_font, fill=title_color)
    if body:
        lines = wrap(draw, body, body_font, x2 - x1 - 48)
        y = y1 + 66
        for line in lines[:5]:
            draw.text((x1 + 24, y), line, font=body_font, fill="#334155")
            y += 32


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = "#334155",
    width: int = 5,
    label: str | None = None,
) -> None:
    draw.line((start, end), fill=color, width=width)
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    pts = [
        end,
        (end[0] - size * math.cos(ang - math.pi / 6), end[1] - size * math.sin(ang - math.pi / 6)),
        (end[0] - size * math.cos(ang + math.pi / 6), end[1] - size * math.sin(ang + math.pi / 6)),
    ]
    draw.polygon(pts, fill=color)
    if label:
        f = pil_font(22, bold=True)
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2
        bbox = draw.textbbox((0, 0), label, font=f)
        pad = 8
        draw.rounded_rectangle(
            (mx - (bbox[2] - bbox[0]) / 2 - pad, my - 19, mx + (bbox[2] - bbox[0]) / 2 + pad, my + 19),
            radius=10,
            fill="white",
            outline=BORDER,
        )
        draw.text((mx - (bbox[2] - bbox[0]) / 2, my - 14), label, font=f, fill=color)


def base_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    im = Image.new("RGB", (1600, 950), "white")
    draw = ImageDraw.Draw(im)
    draw.rectangle((0, 0, 1600, 110), fill=LIGHT)
    draw.text((58, 35), title, font=pil_font(42, bold=True), fill=NAVY)
    draw.line((58, 105, 1542, 105), fill=BORDER, width=2)
    return im, draw


def save_diagram(im: Image.Image, name: str) -> Path:
    path = ASSETS / name
    im.save(path, "PNG", optimize=True)
    return path


def make_diagrams() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    diagrams: dict[str, Path] = {}

    im, d = base_canvas("Kiến trúc tổng quan hệ thống P2P Chat")
    box(d, (610, 165, 990, 365), "Bootstrap Server", "Express API\nPeer discovery\nGroup metadata\nOffline queue", MINT, TEAL)
    box(d, (95, 570, 430, 780), "Peer A - Alice", "Web 3101\nTCP server 5101\nTCP client", "#eff6ff", BLUE)
    box(d, (635, 570, 970, 780), "Peer B - Bob", "Web 3102\nTCP server 5102\nTCP client", "#eff6ff", BLUE)
    box(d, (1175, 570, 1510, 780), "Peer C - Carol", "Web 3103\nTCP server 5103\nTCP client", "#eff6ff", BLUE)
    box(d, (635, 205, 965, 330), "MySQL / Memory Store", "peer, group, ACK, log, offline message", "#fff7ed", ORANGE)
    arrow(d, (265, 570), (610, 330), TEAL, label="register / heartbeat")
    arrow(d, (800, 570), (800, 365), TEAL, label="discovery")
    arrow(d, (1345, 570), (990, 330), TEAL, label="metadata")
    arrow(d, (430, 675), (635, 675), BLUE, label="TCP direct")
    arrow(d, (970, 675), (1175, 675), BLUE, label="TCP group/file")
    arrow(d, (430, 735), (1175, 735), BLUE, label="broadcast")
    diagrams["architecture"] = save_diagram(im, "diagram-architecture.png")

    im, d = base_canvas("Sơ đồ triển khai local và mapping port")
    box(d, (120, 180, 470, 360), "Bootstrap", "http://127.0.0.1:3000\nREST API\nLauncher UI", "#ecfeff", TEAL)
    for i, (name, web, tcp, x) in enumerate([("Alice", 3101, 5101, 560), ("Bob", 3102, 5102, 920), ("Carol", 3103, 5103, 1280)]):
        box(d, (x, 180, x + 260, 360), name, f"Web UI: {web}\nTCP: {tcp}\nPeer process", "#f8fafc", BLUE)
        arrow(d, (x + 130, 360), (x + 130, 560), BLUE, label="Node.js")
        box(d, (x - 5, 560, x + 265, 740), f"Process {i+1}", "Express + Socket.IO\nTCP server\nPeerRuntime", "#eff6ff", BLUE)
    arrow(d, (470, 270), (560, 270), TEAL, label="spawn")
    arrow(d, (820, 270), (920, 270), TEAL, label="register")
    arrow(d, (1180, 270), (1280, 270), TEAL, label="register")
    diagrams["deployment"] = save_diagram(im, "diagram-deployment.png")

    im, d = base_canvas("Các thành phần trong Bootstrap Server")
    labels = [
        ("Launcher UI", "start/stop peer process", 90, 190, "#ecfeff", TEAL),
        ("Peer Registry", "register, unregister, heartbeat", 460, 190, "#f0fdf4", GREEN),
        ("Group Service", "group, members, ownership", 830, 190, "#eff6ff", BLUE),
        ("Offline Queue", "store-and-forward khi peer offline", 1200, 190, "#fff7ed", ORANGE),
        ("Message Store", "direct, group, ACK", 275, 560, "#f8fafc", NAVY),
        ("System Logs", "audit runtime events", 650, 560, "#f8fafc", NAVY),
        ("Storage Adapter", "MySQL hoặc memory fallback", 1025, 560, "#fee2e2", RED),
    ]
    for title, body, x, y, fill, color in labels:
        box(d, (x, y, x + 300, y + 170), title, body, fill, color)
    for x in [240, 610, 980, 1350]:
        arrow(d, (x, 360), (800, 560), GRAY)
    arrow(d, (575, 645), (650, 645), GRAY)
    arrow(d, (950, 645), (1025, 645), GRAY)
    diagrams["bootstrap_components"] = save_diagram(im, "diagram-bootstrap-components.png")

    im, d = base_canvas("Các thành phần trong Peer App")
    box(d, (80, 170, 390, 330), "Web UI", "EJS + CSS\nDirect, Group, File\nMessage history", "#eff6ff", BLUE)
    box(d, (485, 170, 800, 330), "Peer API", "REST endpoint nội bộ\nvalidate request\nmulter upload", "#ecfeff", TEAL)
    box(d, (895, 170, 1245, 330), "PeerRuntime", "orchestrate TCP\nsync bootstrap\nstate + stats", "#f0fdf4", GREEN)
    box(d, (1290, 170, 1530, 330), "Socket.IO", "push realtime UI", "#f8fafc", NAVY)
    box(d, (280, 560, 565, 730), "TCP Client", "send frame\nretry + timeout\nwait ACK", "#fff7ed", ORANGE)
    box(d, (660, 560, 945, 730), "TCP Server", "parse frame\nhandle payload\nreturn ACK", "#fff7ed", ORANGE)
    box(d, (1040, 560, 1325, 730), "Bootstrap Client", "HTTP register\npeers, groups\noffline queue", "#ecfeff", TEAL)
    arrow(d, (390, 250), (485, 250), BLUE)
    arrow(d, (800, 250), (895, 250), GREEN)
    arrow(d, (1245, 250), (1290, 250), NAVY)
    arrow(d, (1010, 330), (430, 560), ORANGE, label="outbound")
    arrow(d, (1065, 330), (805, 560), ORANGE, label="inbound")
    arrow(d, (1110, 330), (1182, 560), TEAL, label="HTTP")
    diagrams["peer_components"] = save_diagram(im, "diagram-peer-components.png")

    im, d = base_canvas("Luồng direct message và ACK")
    lanes = [("Peer A", 160), ("TCP Network", 560), ("Peer B", 960), ("Bootstrap", 1320)]
    for name, x in lanes:
        draw_x = x
        d.text((draw_x - 70, 150), name, font=pil_font(28, bold=True), fill=NAVY)
        d.line((draw_x, 200, draw_x, 820), fill=BORDER, width=4)
    steps = [
        (160, 260, 960, 260, "1. gửi JSON frame"),
        (960, 340, 160, 340, "2. ACK received"),
        (960, 430, 1320, 430, "3. lưu message"),
        (160, 520, 1320, 520, "4. lưu ACK delivered"),
        (1320, 610, 160, 610, "5. UI cập nhật history"),
    ]
    for sx, sy, ex, ey, label in steps:
        arrow(d, (sx, sy), (ex, ey), BLUE if sx < ex else GREEN, label=label)
    box(d, (90, 710, 1510, 850), "Điểm chính", "Nội dung chat đi trực tiếp giữa TCP client của Peer A và TCP server của Peer B. Bootstrap chỉ lưu metadata và lịch sử.", "#f8fafc", NAVY)
    diagrams["direct_flow"] = save_diagram(im, "diagram-direct-flow.png")

    im, d = base_canvas("Luồng group message")
    box(d, (95, 170, 450, 340), "Peer A", "lấy danh sách member\nmã hóa payload\nlặp qua từng peer", "#eff6ff", BLUE)
    box(d, (625, 150, 975, 310), "Bootstrap", "lưu group metadata\nowner + members", MINT, TEAL)
    box(d, (1150, 170, 1505, 340), "Group Members", "peer-b\npeer-c\nmỗi peer ACK riêng", "#f0fdf4", GREEN)
    box(d, (330, 570, 650, 760), "TCP send #1", "Alice -> Bob\nACK Bob", "#fff7ed", ORANGE)
    box(d, (760, 570, 1080, 760), "TCP send #2", "Alice -> Carol\nACK Carol", "#fff7ed", ORANGE)
    arrow(d, (450, 255), (625, 230), TEAL, label="GET/POST group")
    arrow(d, (975, 230), (1150, 255), TEAL, label="members")
    arrow(d, (270, 340), (490, 570), BLUE, label="for each")
    arrow(d, (270, 340), (920, 570), BLUE, label="for each")
    arrow(d, (650, 665), (760, 665), GREEN, label="aggregate status")
    diagrams["group_flow"] = save_diagram(im, "diagram-group-flow.png")

    im, d = base_canvas("Luồng broadcast tới các peer online")
    box(d, (110, 180, 470, 350), "Peer A", "broadcast(content)\ngetOnlinePeers()", "#eff6ff", BLUE)
    for i, (name, x) in enumerate([("Peer B", 730), ("Peer C", 1130)]):
        box(d, (x, 180, x + 270, 350), name, "nhận TCP frame\ntrả ACK\nhiển thị realtime", "#f0fdf4", GREEN)
        arrow(d, (470, 265), (x, 265), BLUE, label=f"send #{i+1}")
    box(d, (555, 570, 1045, 760), "Kết quả", "Một message logic được nhân bản thành nhiều TCP delivery. Nếu một peer lỗi, trạng thái tổng hợp là partial hoặc queued.", "#fff7ed", ORANGE)
    arrow(d, (865, 350), (865, 570), ORANGE, label="delivery results")
    diagrams["broadcast_flow"] = save_diagram(im, "diagram-broadcast-flow.png")

    im, d = base_canvas("Store-and-forward khi peer offline")
    box(d, (90, 170, 410, 330), "Peer A", "gửi tới Carol\nretry 3 lần\nnhận ECONNREFUSED", "#eff6ff", BLUE)
    box(d, (645, 170, 960, 330), "Bootstrap Queue", "offline_messages\nstatus=pending", "#fff7ed", ORANGE)
    box(d, (1190, 170, 1510, 330), "Peer C offline", "TCP server tắt\nkhông trả ACK", "#fee2e2", RED)
    box(d, (1190, 585, 1510, 760), "Peer C online lại", "poll queue\nhandle payload\nmark delivered", "#f0fdf4", GREEN)
    arrow(d, (410, 250), (1190, 250), RED, label="TCP fail")
    arrow(d, (410, 310), (645, 250), ORANGE, label="store")
    arrow(d, (1345, 585), (960, 330), TEAL, label="poll")
    arrow(d, (960, 250), (1190, 665), GREEN, label="deliver")
    diagrams["offline_flow"] = save_diagram(im, "diagram-offline-flow.png")

    im, d = base_canvas("Máy trạng thái gửi tin và ACK")
    states = [
        ("sending", 160, 220, BLUE),
        ("retrying", 520, 220, ORANGE),
        ("delivered", 900, 220, GREEN),
        ("failed", 520, 600, RED),
        ("queued_offline", 900, 600, ORANGE),
    ]
    for title, x, y, color in states:
        box(d, (x, y, x + 250, y + 140), title, "", "#f8fafc", color)
        centered_text(d, (x, y + 40, x + 250, y + 125), title, pil_font(30, bold=True), color)
    arrow(d, (410, 290), (520, 290), ORANGE, label="timeout")
    arrow(d, (770, 290), (900, 290), GREEN, label="ACK")
    arrow(d, (645, 360), (645, 600), RED, label="hết retry")
    arrow(d, (770, 665), (900, 665), ORANGE, label="queue=true")
    arrow(d, (285, 360), (900, 290), GREEN, label="ACK ngay")
    diagrams["state_machine"] = save_diagram(im, "diagram-state-machine.png")

    im, d = base_canvas("ERD cơ sở dữ liệu")
    tables = [
        ("users", "id PK\nusername\ncreated_at", 70, 165),
        ("peers", "peer_id PK\nusername\nhost, ports\nstatus", 390, 165),
        ("chat_groups", "group_id PK\nname\nowner_peer_id", 720, 165),
        ("group_members", "group_id FK\npeer_id FK\njoined_at", 1050, 165),
        ("direct_messages", "message_id PK\nfrom_peer_id\nto_peer_id\nstatus", 220, 540),
        ("group_messages", "message_id PK\ngroup_id\nfrom_peer_id\nstatus", 570, 540),
        ("offline_messages", "id PK\ntarget_peer_id\nmessage_payload\nstatus", 920, 540),
        ("message_acks", "message_id\nfrom_peer_id\nto_peer_id\nattempts", 1270, 540),
    ]
    coords = {}
    for name, body, x, y in tables:
        box(d, (x, y, x + 250, y + 190), name, body, "#f8fafc", NAVY)
        coords[name] = (x, y, x + 250, y + 190)
    arrow(d, (320, 255), (390, 255), GRAY)
    arrow(d, (970, 255), (1050, 255), GRAY)
    arrow(d, (845, 355), (695, 540), GRAY)
    arrow(d, (515, 355), (345, 540), GRAY)
    arrow(d, (1175, 540), (1395, 540), GRAY)
    arrow(d, (820, 635), (920, 635), GRAY)
    diagrams["erd"] = save_diagram(im, "diagram-erd.png")

    im, d = base_canvas("Luồng mã hóa AES-256-GCM")
    stages = [
        ("Plain text", "nội dung người dùng nhập", 90, "#eff6ff", BLUE),
        ("Encrypt", "AES-256-GCM\nsecret + iv", 410, "#f0fdf4", GREEN),
        ("Wire payload", "iv, tag, data\nkhông gửi plain text", 730, "#fff7ed", ORANGE),
        ("TCP", "newline-delimited JSON frame", 1050, "#f8fafc", NAVY),
        ("Decrypt", "verify tag\nkhôi phục content", 1320, "#f0fdf4", GREEN),
    ]
    for title, body, x, fill, color in stages:
        box(d, (x, 250, x + 240, 440), title, body, fill, color)
    for x1, x2 in [(330, 410), (650, 730), (970, 1050), (1290, 1320)]:
        arrow(d, (x1, 345), (x2, 345), TEAL)
    box(d, (280, 620, 1320, 780), "Ý nghĩa bảo mật", "Nội dung tin nhắn được bảo vệ tính bí mật và toàn vẹn trên đường truyền TCP. Trong phạm vi đồ án, các peer dùng shared secret cấu hình bằng biến môi trường.", "#f8fafc", NAVY)
    diagrams["encryption"] = save_diagram(im, "diagram-encryption.png")

    im, d = base_canvas("Luồng truyền file")
    box(d, (90, 180, 390, 360), "Web UI", "chọn file\nsubmit multipart", "#eff6ff", BLUE)
    box(d, (500, 180, 800, 360), "Multer", "memory storage\ncheck maxFileBytes", "#ecfeff", TEAL)
    box(d, (910, 180, 1210, 360), "TCP Payload", "base64 data\nmetadata file", "#fff7ed", ORANGE)
    box(d, (1320, 180, 1540, 360), "Peer nhận", "ghi received_files\nlưu metadata", "#f0fdf4", GREEN)
    arrow(d, (390, 270), (500, 270), BLUE)
    arrow(d, (800, 270), (910, 270), ORANGE)
    arrow(d, (1210, 270), (1320, 270), GREEN)
    box(d, (380, 590, 1220, 760), "Giới hạn hiện tại", "File được đóng gói base64 nên phù hợp file nhỏ. Hướng phát triển là streaming hoặc chia chunk để truyền file lớn hiệu quả hơn.", "#f8fafc", NAVY)
    diagrams["file_flow"] = save_diagram(im, "diagram-file-flow.png")

    im, d = base_canvas("Mô phỏng churn trong hệ phân tán")
    box(d, (130, 260, 420, 430), "Online", "register\nheartbeat\nTCP server on", "#f0fdf4", GREEN)
    box(d, (650, 260, 950, 430), "Leave", "unregister\nstop TCP server\nstatus offline", "#fee2e2", RED)
    box(d, (1180, 260, 1470, 430), "Rejoin", "start TCP server\nregister lại\nsync state", "#eff6ff", BLUE)
    arrow(d, (420, 345), (650, 345), RED, label="timer")
    arrow(d, (950, 345), (1180, 345), BLUE, label="timer")
    arrow(d, (1325, 430), (275, 430), TEAL, label="loop")
    box(d, (365, 620, 1240, 770), "Mục tiêu kiểm thử", "Quan sát hệ thống khi peer rời mạng, timeout, offline queue và cập nhật danh sách peer online/offline.", "#f8fafc", NAVY)
    diagrams["churn"] = save_diagram(im, "diagram-churn.png")

    im, d = base_canvas("Bản đồ API chính")
    box(d, (75, 170, 710, 760), "Bootstrap API", "GET /health\nPOST /api/register\nPOST /api/heartbeat\nGET /api/peers\nPOST /api/groups\nPOST /api/offline-messages\nGET /api/messages/:peerId\nPOST /api/acks\nGET /api/launcher", "#ecfeff", TEAL)
    box(d, (890, 170, 1525, 760), "Peer API", "GET /api/me\nGET /api/peers\nPOST /api/messages/direct\nPOST /api/messages/group\nPOST /api/broadcast\nPOST /api/groups\nPOST /api/files\nPOST /api/churn/start\nPOST /api/sync", "#eff6ff", BLUE)
    arrow(d, (710, 460), (890, 460), NAVY, label="Web UI và runtime gọi HTTP")
    diagrams["api_map"] = save_diagram(im, "diagram-api-map.png")

    im, d = base_canvas("Mô hình lưu trữ MySQL và memory fallback")
    box(d, (130, 200, 500, 380), "Ứng dụng", "Bootstrap server\nstore abstraction", "#eff6ff", BLUE)
    box(d, (680, 150, 1050, 330), "MySQL", "production/demo đầy đủ\nschema.sql", "#ecfeff", TEAL)
    box(d, (680, 480, 1050, 660), "Memory Store", "fallback khi DB chưa chạy\nphù hợp luyện demo", "#fff7ed", ORANGE)
    box(d, (1190, 310, 1500, 500), "Dữ liệu", "peers\ngroups\nmessages\nACK\noffline queue\nlogs", "#f8fafc", NAVY)
    arrow(d, (500, 290), (680, 240), TEAL, label="DB_ENABLED")
    arrow(d, (500, 320), (680, 570), ORANGE, label="fallback")
    arrow(d, (1050, 240), (1190, 385), NAVY)
    arrow(d, (1050, 570), (1190, 425), NAVY)
    diagrams["storage"] = save_diagram(im, "diagram-storage.png")

    im, d = base_canvas("Ma trận kiểm thử chức năng")
    cols = [90, 390, 720, 1050, 1380]
    rows = [170, 320, 470, 620]
    headers = ["Discovery", "Direct", "Group", "Offline", "File"]
    bodies = ["peer online", "ACK delivered", "mỗi member", "queued + delivered", "save metadata"]
    for i, x in enumerate(cols):
        box(d, (x, 210, x + 220, 350), headers[i], bodies[i], "#f8fafc", [TEAL, BLUE, GREEN, ORANGE, NAVY][i])
    for i in range(4):
        arrow(d, (cols[i] + 220, 280), (cols[i + 1], 280), GRAY)
    box(d, (210, 575, 1390, 760), "Kết quả mong đợi", "Các test thể hiện đúng bản chất P2P: message chính đi qua TCP peer-to-peer, còn bootstrap hỗ trợ discovery, metadata và store-and-forward.", "#f0fdf4", GREEN)
    diagrams["testing"] = save_diagram(im, "diagram-testing.png")

    return diagrams


def prepare_screenshots() -> dict[str, Path]:
    shots = {
        "bootstrap": "report-bootstrap-launcher.png",
        "peer_a_direct": "report-peer-a-direct.png",
        "peer_a_group": "report-peer-a-group.png",
        "peer_a_broadcast": "report-peer-a-broadcast.png",
        "peer_a_file": "report-peer-a-file.png",
        "peer_b_received": "report-peer-b-received.png",
        "peer_b_file": "report-peer-b-file-received.png",
        "peer_a_offline": "report-peer-a-offline-queue.png",
        "peer_c_delivered": "report-peer-c-offline-delivered.png",
    }
    out: dict[str, Path] = {}
    for key, filename in shots.items():
        src = ROOT / filename
        dst = ASSETS / filename
        if src.exists():
            with Image.open(src) as im:
                im = im.convert("RGB")
                im.thumbnail((1700, 1200), Image.Resampling.LANCZOS)
                bordered = ImageOps.expand(im, border=3, fill="#e2e8f0")
                bordered.save(dst, "PNG", optimize=True)
            out[key] = dst
    return out


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run = paragraph.add_run("1")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, NAVY),
        ("Heading 2", 14, TEAL),
        ("Heading 3", 13, "334155"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = True
    for s in doc.sections:
        add_page_number(s)
    configure_styles(doc)


def add_para(doc: Document, text: str = "", bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(13)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(13)


FIGURES: list[str] = []
TABLES: list[str] = []


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 15.6) -> None:
    if not path.exists():
        add_para(doc, f"[Thiếu hình: {caption}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    r = cap.add_run(f"Hình {len(FIGURES) + 1}. {caption}")
    r.italic = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)
    FIGURES.append(caption)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"Bảng {len(TABLES) + 1}. {caption}")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], NAVY)
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    TABLES.append(caption)


def add_code_block(doc: Document, code: str, caption: str | None = None) -> None:
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12.5)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "#f1f5f9")
    cell.text = ""
    for line in code.strip("\n").splitlines():
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    for text, size, bold, space in [
        ("TRƯỜNG ĐẠI HỌC", 15, True, 0),
        ("KHOA CÔNG NGHỆ THÔNG TIN", 14, True, 70),
        ("BÁO CÁO ĐỒ ÁN TỐT NGHIỆP", 18, True, 30),
        ("ĐỀ TÀI", 15, True, 8),
        ("XÂY DỰNG HỆ THỐNG CHAT NGANG HÀNG P2P", 20, True, 12),
        ("Peer-to-Peer Chat System", 15, False, 80),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(size)
        r.bold = bold
    info = [
        ("Môn học", "Các hệ thống phân tán"),
        ("Chủ đề", "Chủ đề 3 - Phát triển hệ thống chat ngang hàng P2P"),
        ("Công nghệ", "Node.js, Express, TCP Socket, Socket.IO, MySQL"),
        ("Sinh viên thực hiện", "........................................................"),
        ("Mã sinh viên", "........................................................"),
        ("Giảng viên hướng dẫn", "........................................................"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(info):
        set_cell_text(table.cell(i, 0), k, bold=True)
        set_cell_text(table.cell(i, 1), v)
        set_cell_shading(table.cell(i, 0), "#f1f5f9")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("TP. Hồ Chí Minh, tháng 05 năm 2026")
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("PHIẾU GIAO NHIỆM VỤ ĐỒ ÁN", level=1)
    add_para(doc, "Tên đề tài: Xây dựng hệ thống chat ngang hàng P2P phục vụ minh họa các nguyên lý của hệ thống phân tán.", bold=True)
    add_table(
        doc,
        "Thông tin nhiệm vụ và phạm vi báo cáo",
        ["Nội dung", "Mô tả"],
        [
            ["Mục tiêu", "Xây dựng hệ thống chat P2P trong đó mỗi peer vừa gửi vừa nhận tin qua TCP socket."],
            ["Phạm vi", "Chạy local hoặc LAN, hỗ trợ nhiều peer, discovery qua bootstrap, lưu metadata bằng MySQL hoặc memory fallback."],
            ["Sản phẩm", "Source code, hướng dẫn chạy, báo cáo thiết kế, ảnh demo và kịch bản kiểm thử."],
            ["Ràng buộc", "Bootstrap không chuyển tiếp luồng chat chính; tin nhắn realtime đi trực tiếp giữa các peer."],
        ],
        widths=[4.0, 11.5],
    )
    add_para(doc, "Các thông tin sinh viên, lớp và giảng viên hướng dẫn được để dạng dòng chấm để người dùng điền theo biểu mẫu chính thức của khoa.")
    doc.add_page_break()

    doc.add_heading("LỜI CẢM ƠN", level=1)
    for t in [
        "Em xin gửi lời cảm ơn tới giảng viên phụ trách học phần Các hệ thống phân tán đã định hướng đề tài, cung cấp kiến thức nền tảng về mô hình peer-to-peer, giao tiếp tiến trình, đồng bộ trạng thái và xử lý lỗi trong môi trường phân tán.",
        "Trong quá trình thực hiện, đề tài giúp em củng cố kỹ năng thiết kế kiến trúc, tổ chức source code Node.js, xây dựng giao diện demo, kiểm thử các tình huống online/offline và trình bày hệ thống theo cách có thể bảo vệ trước hội đồng.",
        "Báo cáo này tổng hợp quá trình phân tích, thiết kế, cài đặt và đánh giá hệ thống P2P Chat. Các nội dung kỹ thuật, hình ảnh demo và sơ đồ luồng được xây dựng trực tiếp từ project trong thư mục mã nguồn.",
    ]:
        add_para(doc, t)
    doc.add_page_break()

    doc.add_heading("TÓM TẮT", level=1)
    for t in [
        "Đề tài xây dựng một hệ thống chat ngang hàng P2P cho phép nhiều peer trao đổi tin nhắn trực tiếp qua TCP socket. Mỗi peer là một process Node.js độc lập, có Web UI để người dùng thao tác, TCP server để nhận dữ liệu và TCP client để gửi dữ liệu tới peer khác.",
        "Bootstrap server đóng vai trò hỗ trợ discovery, heartbeat, metadata nhóm, offline queue và launcher phục vụ demo. Điểm quan trọng là bootstrap không nằm trên đường truyền tin nhắn chính. Khi Alice gửi tin cho Bob, dữ liệu được gửi trực tiếp từ TCP client của Alice tới TCP server của Bob, sau đó Bob trả ACK để Alice cập nhật trạng thái delivered.",
        "Hệ thống hỗ trợ direct chat, group chat, broadcast, file transfer, mã hóa AES-256-GCM, retry/timeout, store-and-forward khi peer offline và mô phỏng churn. MySQL được dùng để lưu trạng thái, lịch sử, ACK, group và log; đồng thời có memory fallback để dễ chạy demo khi chưa cấu hình database.",
    ]:
        add_para(doc, t)
    add_table(
        doc,
        "Từ khóa chính của đề tài",
        ["Nhóm từ khóa", "Từ khóa"],
        [
            ["Mô hình", "Peer-to-peer, distributed system, bootstrap discovery"],
            ["Giao tiếp", "TCP socket, JSON line, ACK, retry, timeout"],
            ["Dữ liệu", "MySQL, offline queue, message history, file metadata"],
            ["Bảo mật", "AES-256-GCM, shared secret, integrity tag"],
        ],
        widths=[4.2, 11.0],
    )
    doc.add_page_break()

    doc.add_heading("ABSTRACT", level=1)
    for t in [
        "This project implements a peer-to-peer chat system in which each peer acts as both a client and a server. The primary chat path is built on direct TCP socket communication between peers, while a bootstrap server is used for peer discovery, heartbeat tracking, group metadata and offline message storage.",
        "The system includes direct messaging, group messaging, network broadcast, file transfer, AES-256-GCM encryption, ACK-based delivery tracking, retry and timeout handling, store-and-forward delivery for offline peers and churn simulation. A realtime Web UI is provided for demonstration and observation.",
        "The report documents requirements, architecture, protocol design, database design, implementation details, testing scenarios, demo screenshots and future improvements.",
    ]:
        add_para(doc, t)
    doc.add_page_break()

    doc.add_heading("MỤC LỤC", level=1)
    toc = [
        "Chương 1. Tổng quan đề tài",
        "Chương 2. Cơ sở lý thuyết và công nghệ sử dụng",
        "Chương 3. Phân tích yêu cầu hệ thống",
        "Chương 4. Thiết kế kiến trúc và luồng xử lý",
        "Chương 5. Thiết kế dữ liệu, giao thức và bảo mật",
        "Chương 6. Cài đặt, triển khai và giao diện demo",
        "Chương 7. Kiểm thử, đánh giá và kết quả thực nghiệm",
        "Chương 8. Kết luận, hạn chế và hướng phát triển",
        "Tài liệu tham khảo",
        "Phụ lục A. Hướng dẫn chạy hệ thống",
        "Phụ lục B. API và cấu hình",
        "Phụ lục C. Kịch bản bảo vệ",
    ]
    for item in toc:
        add_para(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    doc.add_heading("DANH MỤC HÌNH", level=1)
    for i, fig in enumerate(PREDECLARED_FIGURES, start=1):
        add_para(doc, f"Hình {i}. {fig}", align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    doc.add_heading("DANH MỤC BẢNG", level=1)
    for i, table in enumerate(PREDECLARED_TABLES, start=1):
        add_para(doc, f"Bảng {i}. {table}", align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    doc.add_heading("DANH MỤC TỪ VIẾT TẮT", level=1)
    add_table(
        doc,
        "Danh mục từ viết tắt",
        ["Từ viết tắt", "Ý nghĩa"],
        [
            ["P2P", "Peer-to-Peer, mô hình ngang hàng"],
            ["TCP", "Transmission Control Protocol"],
            ["ACK", "Acknowledgement, gói xác nhận đã nhận message"],
            ["UI", "User Interface"],
            ["API", "Application Programming Interface"],
            ["DB", "Database"],
            ["TTL", "Time To Live"],
            ["GCM", "Galois/Counter Mode trong AES"],
        ],
        widths=[4.0, 11.5],
    )
    doc.add_page_break()


def chapter_1(doc: Document) -> None:
    doc.add_heading("CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", level=1)
    doc.add_heading("1.1. Lý do chọn đề tài", level=2)
    paragraphs = [
        "Trong các hệ thống phân tán, mô hình ngang hàng P2P là một mô hình quan trọng vì mỗi node không chỉ tiêu thụ dịch vụ mà còn có thể cung cấp dịch vụ cho node khác. Mô hình này thường xuất hiện trong chia sẻ file, blockchain, truyền thông thời gian thực và nhiều ứng dụng cần giảm phụ thuộc vào máy chủ trung tâm.",
        "Đề tài chat ngang hàng P2P phù hợp để minh họa nhiều vấn đề cốt lõi của hệ phân tán: phát hiện node, giao tiếp qua mạng, xử lý lỗi kết nối, đồng bộ trạng thái, lưu lịch sử, quan sát runtime và kiểm thử khi một node rời mạng.",
        "Project trong workspace triển khai một mạng chat gồm nhiều peer độc lập. Mỗi peer có Web UI, TCP server và TCP client riêng. Bootstrap server chỉ giữ vai trò danh bạ và metadata, không chuyển tiếp tin nhắn chính. Vì vậy hệ thống thể hiện đúng điểm khác biệt giữa P2P chat và client-server chat truyền thống.",
    ]
    for t in paragraphs:
        add_para(doc, t)
    doc.add_heading("1.2. Mục tiêu nghiên cứu và xây dựng", level=2)
    add_bullets(
        doc,
        [
            "Xây dựng nhiều peer có thể chạy đồng thời trên các port khác nhau.",
            "Cho phép peer đăng ký, heartbeat và lấy danh sách peer online qua bootstrap.",
            "Gửi direct message, group message và broadcast bằng TCP socket trực tiếp giữa các peer.",
            "Có ACK, retry, timeout và offline queue để xử lý tình huống peer đích không nhận được message.",
            "Cung cấp Web UI realtime để demo trạng thái, lịch sử message, group, log và thống kê delivery.",
            "Lưu metadata, lịch sử, ACK, file transfer và system log bằng MySQL hoặc memory fallback.",
        ],
    )
    doc.add_heading("1.3. Phạm vi và đối tượng sử dụng", level=2)
    add_para(doc, "Hệ thống được thiết kế để chạy trong môi trường local hoặc mạng LAN. Mỗi peer được cấu hình bằng biến môi trường gồm peer id, username, host, TCP port và Web port. Người dùng thao tác qua trình duyệt, còn runtime xử lý mạng chạy trong Node.js.")
    add_para(doc, "Đối tượng sử dụng chính là sinh viên hoặc người đánh giá đồ án muốn quan sát một hệ thống phân tán nhỏ nhưng có đầy đủ các tình huống: peer join, peer discovery, gửi tin trực tiếp, gửi nhóm, broadcast, file transfer, peer offline và giao lại message.")
    add_table(
        doc,
        "Phạm vi chức năng của đề tài",
        ["Nhóm chức năng", "Mô tả"],
        [
            ["Cơ bản", "Peer discovery, direct chat, group chat, trạng thái online/offline, Web UI"],
            ["Độ tin cậy", "ACK, timeout, retry, offline queue, store-and-forward"],
            ["Nâng cao", "Broadcast, file transfer, encryption, churn simulation, launcher UI"],
            ["Ngoài phạm vi", "NAT traversal Internet, phân quyền user đầy đủ, truyền file lớn bằng streaming"],
        ],
        widths=[4.5, 11],
    )
    doc.add_page_break()
    doc.add_heading("1.4. Cấu trúc báo cáo", level=2)
    add_para(doc, "Báo cáo được tổ chức theo cấu trúc của một báo cáo đồ án tốt nghiệp: phần đầu trình bày nhiệm vụ, tóm tắt và danh mục; các chương chính trình bày nền tảng, yêu cầu, thiết kế, cài đặt, kiểm thử và đánh giá; phần cuối là tài liệu tham khảo và phụ lục hướng dẫn chạy.")
    add_numbered(
        doc,
        [
            "Chương 1 giới thiệu tổng quan đề tài, mục tiêu và phạm vi.",
            "Chương 2 trình bày cơ sở lý thuyết và công nghệ.",
            "Chương 3 phân tích yêu cầu chức năng, phi chức năng và use case.",
            "Chương 4 mô tả kiến trúc, thành phần và các luồng xử lý chính.",
            "Chương 5 thiết kế database, giao thức TCP và bảo mật.",
            "Chương 6 mô tả cài đặt, triển khai và giao diện demo.",
            "Chương 7 trình bày kiểm thử, ảnh demo và đánh giá kết quả.",
            "Chương 8 kết luận, nêu hạn chế và hướng phát triển.",
        ],
    )


def chapter_2(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", level=1)
    doc.add_heading("2.1. Mô hình peer-to-peer", level=2)
    for t in [
        "Mô hình peer-to-peer là mô hình trong đó các node có quyền và trách nhiệm tương đối ngang nhau. Một peer có thể gửi yêu cầu tới peer khác, đồng thời cũng mở cổng để phục vụ yêu cầu từ peer khác. Khác với mô hình client-server, P2P giảm phụ thuộc vào một server trung tâm trên đường truyền dữ liệu chính.",
        "Trong project này, mỗi peer là một process Node.js độc lập. Peer tự mở TCP server để nhận message và dùng TCP client để gửi message. Bootstrap chỉ hỗ trợ discovery và metadata. Khi đánh giá hệ thống, cần phân biệt rõ đường truyền control plane và data plane: control plane đi qua bootstrap, data plane chat đi trực tiếp qua TCP.",
    ]:
        add_para(doc, t)
    add_figure(doc, diagrams["architecture"], "Kiến trúc tổng quan giữa bootstrap và các peer")
    doc.add_heading("2.2. TCP socket và JSON line protocol", level=2)
    add_para(doc, "TCP cung cấp kênh truyền tin cậy theo stream. Vì TCP không tự chia message theo ranh giới ứng dụng, project dùng định dạng JSON line: mỗi payload JSON kết thúc bằng ký tự xuống dòng. TCP server nhận buffer, tách theo newline và parse từng frame.")
    add_code_block(
        doc,
        """
{
  "type": "direct_message",
  "messageId": "uuid",
  "fromPeerId": "peer-a",
  "toPeerId": "peer-b",
  "encrypted": true,
  "createdAt": "2026-05-20T10:00:00.000Z"
}
""",
        "Ví dụ 1. Payload direct message dạng JSON",
    )
    doc.add_heading("2.3. Công nghệ chính", level=2)
    add_table(
        doc,
        "Stack công nghệ sử dụng trong project",
        ["Thành phần", "Công nghệ", "Vai trò"],
        [
            ["Runtime", "Node.js", "Chạy bootstrap server và peer process, xử lý I/O bất đồng bộ."],
            ["Web backend", "Express", "Cung cấp REST API và render Web UI."],
            ["Realtime UI", "Socket.IO", "Đẩy message, log và trạng thái lên trình duyệt."],
            ["P2P networking", "node:net TCP socket", "Gửi và nhận payload trực tiếp giữa peer."],
            ["Database", "MySQL", "Lưu peers, groups, messages, ACK, offline queue và logs."],
            ["Template", "EJS", "Render giao diện server-side đơn giản, dễ demo."],
            ["File upload", "Multer", "Nhận file từ Web UI trước khi đóng gói vào TCP payload."],
            ["Security", "AES-256-GCM", "Mã hóa và kiểm tra toàn vẹn nội dung message."],
        ],
        widths=[3.4, 4.0, 8.0],
    )
    add_figure(doc, diagrams["deployment"], "Sơ đồ triển khai local với Bootstrap và ba peer demo")
    doc.add_heading("2.4. Đặc trưng hệ thống phân tán được minh họa", level=2)
    add_bullets(
        doc,
        [
            "Nhiều tiến trình độc lập cùng hoạt động và giao tiếp qua mạng.",
            "Trạng thái online/offline của peer thay đổi theo thời gian.",
            "Message delivery phụ thuộc vào network, timeout, ACK và retry.",
            "Bootstrap hỗ trợ discovery nhưng không làm mất bản chất P2P của luồng chat.",
            "Store-and-forward xử lý tình huống một peer rời mạng rồi quay lại.",
            "Churn simulation cho phép quan sát join/leave liên tục.",
        ],
    )


def chapter_3(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG 3. PHÂN TÍCH YÊU CẦU HỆ THỐNG", level=1)
    doc.add_heading("3.1. Yêu cầu chức năng", level=2)
    add_table(
        doc,
        "Yêu cầu chức năng chi tiết",
        ["Mã", "Yêu cầu", "Mức độ"],
        [
            ["F01", "Peer đăng ký với bootstrap khi tham gia mạng.", "Bắt buộc"],
            ["F02", "Bootstrap trả danh sách peer online phục vụ discovery.", "Bắt buộc"],
            ["F03", "Peer gửi direct message tới peer khác bằng TCP.", "Bắt buộc"],
            ["F04", "Peer gửi group message tới từng thành viên trong group.", "Bắt buộc"],
            ["F05", "Peer gửi broadcast tới toàn bộ peer online.", "Nâng cao"],
            ["F06", "Peer nhận ACK và cập nhật trạng thái delivered.", "Bắt buộc"],
            ["F07", "Retry và timeout khi peer đích không phản hồi.", "Bắt buộc"],
            ["F08", "Queue offline message và giao lại khi peer online.", "Nâng cao"],
            ["F09", "Gửi file nhỏ qua TCP payload.", "Nâng cao"],
            ["F10", "Web UI realtime hiển thị peers, messages, groups và logs.", "Bắt buộc"],
        ],
        widths=[1.5, 11.5, 2.6],
    )
    doc.add_heading("3.2. Yêu cầu phi chức năng", level=2)
    add_bullets(
        doc,
        [
            "Tính đúng mô hình P2P: dữ liệu chat chính không được chuyển tiếp qua bootstrap.",
            "Tính quan sát: UI phải thể hiện trạng thái peer, delivery status, log và thống kê.",
            "Tính dễ chạy demo: có launcher để start/stop peer và có memory fallback nếu MySQL chưa sẵn sàng.",
            "Tính mở rộng vừa phải: source code chia module bootstrap, peer, TCP, shared protocol và database store.",
            "Tính bảo mật cơ bản: nội dung message có thể mã hóa trước khi gửi qua TCP.",
        ],
    )
    doc.add_heading("3.3. Tác nhân và ca sử dụng", level=2)
    add_table(
        doc,
        "Tác nhân trong hệ thống",
        ["Tác nhân", "Vai trò"],
        [
            ["Người dùng peer", "Mở Web UI của peer, gửi tin nhắn, tạo nhóm, gửi file, xem lịch sử và log."],
            ["Peer runtime", "Thực hiện đăng ký, heartbeat, gửi/nhận TCP, retry, ACK, sync group và offline queue."],
            ["Bootstrap server", "Quản lý danh bạ peer, metadata group, offline queue, message history và launcher."],
            ["Database store", "Lưu dữ liệu bền vững hoặc dùng memory fallback để demo."],
        ],
        widths=[4, 11.5],
    )
    add_numbered(
        doc,
        [
            "Start hệ thống: người dùng chạy bootstrap, sau đó tạo các peer Alice, Bob và Carol.",
            "Peer discovery: peer gọi bootstrap để lấy danh sách peer đang online.",
            "Direct chat: Alice chọn Bob và gửi message qua TCP.",
            "Group chat: Alice tạo nhóm gồm Bob, Carol và gửi message tới từng thành viên.",
            "Offline delivery: Alice gửi Carol khi Carol offline, message được queue và giao lại khi Carol online.",
            "File transfer: Alice gửi README.md tới Bob, Bob lưu file vào thư mục received_files.",
        ],
    )
    doc.add_heading("3.4. Tiêu chí nghiệm thu", level=2)
    add_table(
        doc,
        "Tiêu chí nghiệm thu đồ án",
        ["Tiêu chí", "Cách kiểm chứng"],
        [
            ["Đúng P2P", "Quan sát code và demo: direct/group/broadcast gọi TCP client tới TCP server của peer nhận."],
            ["Có discovery", "Mỗi peer hiển thị danh sách peer khác online sau khi sync bootstrap."],
            ["Có ACK", "Peer gửi chỉ đánh dấu delivered khi nhận ACK đúng messageId."],
            ["Có offline queue", "Stop Carol, gửi message, thấy failed_queued; start Carol, thấy delivered_from_queue."],
            ["Có giao diện demo", "Bootstrap launcher và Web UI peer hoạt động trên trình duyệt."],
        ],
        widths=[5, 10.5],
    )


def chapter_4(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG 4. THIẾT KẾ KIẾN TRÚC VÀ LUỒNG XỬ LÝ", level=1)
    doc.add_heading("4.1. Kiến trúc tổng thể", level=2)
    add_para(doc, "Kiến trúc hệ thống tách rõ hai nhóm thành phần: bootstrap server và peer app. Bootstrap giữ vai trò control plane, còn peer app thực hiện data plane P2P. Việc tách này giúp project vừa dễ demo, vừa thể hiện đúng yêu cầu không phụ thuộc hoàn toàn vào server trung tâm.")
    add_figure(doc, diagrams["bootstrap_components"], "Các module chính trong Bootstrap Server")
    add_para(doc, "Bootstrap server triển khai bằng Express. Các API quan trọng gồm register, unregister, heartbeat, peers, groups, offline messages, messages, ACK, file transfers, logs và launcher. Store abstraction cho phép dùng MySQL hoặc memory fallback.")
    add_figure(doc, diagrams["peer_components"], "Các module chính trong Peer App")
    add_para(doc, "Peer app gồm Web UI, API nội bộ, PeerRuntime, TCP server, TCP client, BootstrapClient và Socket.IO. PeerRuntime là module điều phối quan trọng nhất vì nó kết nối tất cả luồng: đăng ký, sync peers, gửi tin, nhận tin, lưu lịch sử, thống kê và log.")
    doc.add_heading("4.2. Luồng peer tham gia mạng", level=2)
    add_numbered(
        doc,
        [
            "Peer process khởi động với PEER_ID, USERNAME, TCP_PORT và WEB_PORT.",
            "Peer mở Express Web UI và TCP server.",
            "Peer gọi POST /api/register lên bootstrap.",
            "Bootstrap lưu peer online và trả danh sách peers hiện có.",
            "Peer định kỳ heartbeat, sync peers, sync groups và poll offline queue.",
        ],
    )
    add_para(doc, "Thiết kế này giúp peer tự phục hồi khi bootstrap restart. Nếu heartbeat trả lỗi peer not registered, peer gọi register lại để khôi phục trạng thái trong bootstrap.")
    doc.add_page_break()
    doc.add_heading("4.3. Luồng gửi direct message", level=2)
    add_figure(doc, diagrams["direct_flow"], "Sequence direct message, ACK và lưu lịch sử")
    add_para(doc, "Khi Alice gửi tin cho Bob, runtime tạo payload bằng createWirePayload. Nếu bật encryption, nội dung được mã hóa bằng AES-256-GCM. Sau đó sendTcpPayload mở TCP connection tới host và tcpPort của Bob, gửi frame và đợi ACK. Khi nhận ACK, Alice lưu trạng thái delivered và cập nhật UI.")
    add_heading_text = [
        "Điểm quan trọng trong luồng direct là message không đi qua bootstrap. Bootstrap chỉ được gọi sau đó để lưu message hoặc ACK. Vì vậy bootstrap có thể xem như hệ thống hỗ trợ, không phải server trung chuyển chat.",
        "Nếu Bob không phản hồi, TCP client retry theo cấu hình SEND_RETRIES và TCP_TIMEOUT_MS. Sau khi hết retry, runtime gọi queueOrFail để lưu offline message nếu queueOffline được bật.",
    ]
    for t in add_heading_text:
        add_para(doc, t)
    doc.add_heading("4.4. Luồng group và broadcast", level=2)
    add_figure(doc, diagrams["group_flow"], "Thiết kế group message gửi riêng tới từng thành viên")
    add_para(doc, "Group metadata được lưu ở bootstrap để các peer biết groupId, name, owner và members. Khi gửi group message, peer gửi không chuyển message cho bootstrap phát tán. Thay vào đó, peer gửi duyệt danh sách members rồi gửi cùng một payload tới từng peer bằng TCP.")
    add_figure(doc, diagrams["broadcast_flow"], "Thiết kế broadcast tới toàn bộ peer online")
    add_para(doc, "Broadcast tương tự group message nhưng target được lấy từ danh sách peer online. Trạng thái tổng hợp của broadcast là delivered nếu tất cả peer nhận thành công, partial nếu có peer lỗi và failed_queued nếu toàn bộ target lỗi hoặc bị queue.")
    doc.add_heading("4.4.1. Nhận xét thiết kế group và broadcast", level=3)
    add_bullets(
        doc,
        [
            "Group phù hợp với phạm vi thành viên cố định, do đó bootstrap chỉ cần lưu group metadata và danh sách member.",
            "Broadcast phù hợp thông báo toàn mạng, target được lấy từ danh sách peer online tại thời điểm gửi.",
            "Cả hai luồng đều giữ nguyên nguyên tắc P2P: peer gửi là nơi nhân bản payload và mở TCP connection tới từng peer nhận.",
            "Việc tách ACK theo từng peer giúp UI thể hiện trạng thái partial khi chỉ một phần target nhận thành công.",
        ],
    )
    doc.add_heading("4.5. Luồng store-and-forward", level=2)
    add_figure(doc, diagrams["offline_flow"], "Store-and-forward khi peer đích offline")
    add_para(doc, "Store-and-forward là cơ chế quan trọng để hệ thống chịu lỗi cơ bản. Khi target peer offline hoặc TCP connection thất bại, message được lưu vào offline_messages với status pending. Khi peer đích online, runtime poll queue, xử lý payload như message nhận qua TCP và báo bootstrap đánh dấu delivered.")
    doc.add_heading("4.6. Máy trạng thái delivery", level=2)
    add_figure(doc, diagrams["state_machine"], "Máy trạng thái gửi tin và xử lý ACK")
    add_para(doc, "Các trạng thái sending, delivered, failed, queued_offline và partial được phản ánh trực tiếp lên UI. Nhờ đó người dùng và người bảo vệ có thể quan sát kết quả gửi tin mà không cần đọc log terminal.")


def chapter_5(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG 5. THIẾT KẾ DỮ LIỆU, GIAO THỨC VÀ BẢO MẬT", level=1)
    doc.add_heading("5.1. Thiết kế cơ sở dữ liệu", level=2)
    add_para(doc, "Database không nằm trên đường truyền message chính. Vai trò của database là lưu trạng thái, metadata và lịch sử để hệ thống có thể quan sát, khôi phục và demo rõ ràng. Project cung cấp schema MySQL trong database/schema.sql.")
    add_figure(doc, diagrams["erd"], "ERD các bảng chính của hệ thống")
    add_table(
        doc,
        "Mô tả các bảng dữ liệu chính",
        ["Bảng", "Vai trò"],
        [
            ["users", "Lưu username và display name."],
            ["peers", "Lưu peer id, host, TCP port, Web port, status và last seen."],
            ["peer_status_logs", "Lưu lịch sử thay đổi trạng thái online/offline."],
            ["chat_groups", "Lưu group id, tên nhóm và owner peer."],
            ["group_members", "Lưu quan hệ nhiều-nhiều giữa group và peer."],
            ["direct_messages", "Lưu lịch sử direct và broadcast message."],
            ["group_messages", "Lưu lịch sử message nhóm."],
            ["offline_messages", "Lưu payload chờ giao khi peer đích offline."],
            ["message_acks", "Lưu ACK, số lần gửi và lỗi nếu có."],
            ["file_transfers", "Lưu metadata file transfer."],
            ["system_logs", "Lưu log hệ thống và launcher."],
        ],
        widths=[4.5, 11],
    )
    doc.add_heading("5.2. Thiết kế giao thức TCP", level=2)
    add_para(doc, "Giao thức nội bộ được định nghĩa trong src/shared/protocol.js. Các message type chính gồm direct_message, group_message, broadcast_message, file_transfer, ack và error. Payload được encode thành JSON cộng newline để TCP server tách frame.")
    add_code_block(
        doc,
        """
export const MESSAGE_TYPES = Object.freeze({
  DIRECT: 'direct_message',
  GROUP: 'group_message',
  BROADCAST: 'broadcast_message',
  FILE: 'file_transfer',
  ACK: 'ack',
  ERROR: 'error'
});
""",
        "Ví dụ 2. Các message type trong protocol",
    )
    add_table(
        doc,
        "Các loại payload chính",
        ["Type", "Ý nghĩa", "ACK"],
        [
            ["direct_message", "Tin nhắn từ một peer tới một peer.", "Có"],
            ["group_message", "Tin nhắn logic của group, được gửi riêng tới từng member.", "Có"],
            ["broadcast_message", "Tin nhắn gửi tới tất cả peer online.", "Có"],
            ["file_transfer", "Payload truyền file nhỏ dạng base64.", "Có"],
            ["ack", "Xác nhận peer nhận đã xử lý message.", "Không"],
        ],
        widths=[4, 8.5, 3],
    )
    doc.add_heading("5.3. Bảo mật nội dung message", level=2)
    add_figure(doc, diagrams["encryption"], "Luồng mã hóa và giải mã AES-256-GCM")
    add_para(doc, "Khi ENCRYPTION_ENABLED=true, content không được gửi plain text qua TCP. Runtime gọi maybeEncryptContent để sinh encryption payload gồm algorithm, iv, tag và data. Peer nhận dùng maybeDecryptContent để xác thực tag và khôi phục nội dung.")
    add_bullets(
        doc,
        [
            "Confidentiality: peer trung gian trên mạng không đọc được plain text nếu không có shared secret.",
            "Integrity: GCM tag giúp phát hiện dữ liệu bị sửa đổi.",
            "Limit: đồ án dùng shared secret chung, chưa triển khai key exchange riêng cho từng peer.",
        ],
    )
    doc.add_heading("5.4. Thiết kế file transfer", level=2)
    add_figure(doc, diagrams["file_flow"], "Luồng gửi file qua TCP payload")
    add_para(doc, "File transfer dùng Multer memory storage để nhận file từ form Web UI. Runtime chuyển buffer thành base64 và đóng gói vào payload file_transfer. Peer nhận giải mã base64, ghi file vào received_files/<peerId> và lưu metadata file_transfers.")
    doc.add_heading("5.5. Storage fallback", level=2)
    add_figure(doc, diagrams["storage"], "Cơ chế MySQL và memory fallback")
    add_para(doc, "DB_FALLBACK_MEMORY giúp project chạy được ngay cả khi MySQL chưa bật. Khi bảo vệ, nên bật MySQL thật để chứng minh dữ liệu được lưu bền vững. Khi luyện demo nhanh, memory fallback giúp giảm lỗi môi trường.")


def chapter_6(doc: Document, diagrams: dict[str, Path], shots: dict[str, Path]) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG 6. CÀI ĐẶT, TRIỂN KHAI VÀ GIAO DIỆN DEMO", level=1)
    doc.add_heading("6.1. Cấu trúc thư mục", level=2)
    add_code_block(
        doc,
        """
p2p-chat-system/
  src/
    bootstrap/        server, launcher, Web UI bootstrap
    peer/             peer server, routes, runtime, TCP client/server, Web UI
    shared/           protocol, crypto, HTTP helper
    database/         store MySQL và memory fallback
  database/schema.sql
  docs/
  scripts/check-syntax.mjs
  received_files/
""",
        "Cây thư mục chính của project",
    )
    add_para(doc, "Cách tổ chức source code tách rõ bootstrap, peer, shared utility và database. Điều này giúp mỗi module có trách nhiệm riêng, giảm trộn logic Web UI, networking và storage.")
    add_heading = doc.add_heading("6.2. Cấu hình môi trường", level=2)
    add_code_block(
        doc,
        """
BOOTSTRAP_PORT=3000
BOOTSTRAP_URL=http://127.0.0.1:3000
PEER_ID=peer-a
USERNAME=Alice
PEER_HOST=127.0.0.1
TCP_PORT=5101
WEB_PORT=3101
DB_ENABLED=true
DB_FALLBACK_MEMORY=true
ENCRYPTION_ENABLED=true
P2P_SHARED_SECRET=change-this-secret-before-demo
""",
        "Ví dụ 3. Cấu hình .env mẫu",
    )
    add_table(
        doc,
        "Scripts npm phục vụ chạy và kiểm tra",
        ["Script", "Chức năng"],
        [
            ["npm run bootstrap", "Chạy bootstrap server ở port 3000."],
            ["npm run peer", "Chạy một peer theo biến môi trường."],
            ["npm run peer:a", "Chạy Alice ở Web 3101 và TCP 5101."],
            ["npm run peer:b", "Chạy Bob ở Web 3102 và TCP 5102."],
            ["npm run peer:c", "Chạy Carol ở Web 3103 và TCP 5103."],
            ["npm run check", "Kiểm tra cú pháp 16 file JavaScript bằng node --check."],
        ],
        widths=[5.0, 10.5],
    )
    doc.add_heading("6.3. Giao diện Bootstrap Launcher", level=2)
    add_para(doc, "Bootstrap Launcher là màn hình giúp start peer mới, xem peer registered, mở Web UI của peer và dừng process do launcher tạo. Launcher không làm sai mô hình P2P vì sau khi peer được start, peer vẫn tự mở TCP server và tự giao tiếp với peer khác.")
    add_figure(doc, shots["bootstrap"], "Bootstrap Launcher hiển thị peer đang online")
    doc.add_heading("6.4. Giao diện Peer Web UI", level=2)
    add_para(doc, "Peer Web UI gồm sidebar thông tin peer, thống kê sent/delivered/failed/queued, danh sách peers, các tab Direct, Group, Broadcast, File, message history, group list và system log. Socket.IO cập nhật realtime nên người dùng không cần refresh khi nhận tin.")
    add_figure(doc, shots["peer_a_direct"], "Peer Alice sau khi gửi direct, group và broadcast")
    add_figure(doc, shots["peer_a_group"], "Tab Group trên Peer Alice")
    add_figure(doc, shots["peer_a_broadcast"], "Tab Broadcast trên Peer Alice")
    add_figure(doc, shots["peer_a_file"], "Tab File transfer trên Peer Alice")
    doc.add_heading("6.5. API map", level=2)
    add_figure(doc, diagrams["api_map"], "Bản đồ API Bootstrap và Peer")
    add_para(doc, "Peer API được dùng bởi Web UI cục bộ của từng peer. Bootstrap API được dùng bởi peer runtime để đăng ký, heartbeat, lấy peers, lưu group, lưu history, lưu ACK và xử lý offline queue.")


def chapter_7(doc: Document, diagrams: dict[str, Path], shots: dict[str, Path]) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG 7. KIỂM THỬ, ĐÁNH GIÁ VÀ KẾT QUẢ THỰC NGHIỆM", level=1)
    doc.add_heading("7.1. Phương pháp kiểm thử", level=2)
    add_para(doc, "Kiểm thử tập trung vào các luồng người dùng và các tình huống lỗi phổ biến của hệ phân tán. Môi trường demo gồm bootstrap server, ba peer Alice, Bob và Carol. Các peer chạy trên các port Web/TCP riêng, dùng memory store để seed dữ liệu nhanh và có thể chuyển sang MySQL khi bảo vệ.")
    add_figure(doc, diagrams["testing"], "Ma trận kiểm thử chức năng chính")
    add_table(
        doc,
        "Kịch bản kiểm thử chức năng",
        ["STT", "Kịch bản", "Kết quả mong đợi"],
        [
            ["1", "Start bootstrap", "Bootstrap chạy ở http://127.0.0.1:3000, /health trả ok."],
            ["2", "Start Alice, Bob, Carol", "Mỗi peer có Web UI và TCP server riêng."],
            ["3", "Peer discovery", "Alice thấy Bob và Carol online."],
            ["4", "Direct Alice -> Bob", "Bob nhận realtime, Alice thấy delivered."],
            ["5", "Group Alice -> Bob, Carol", "Bob và Carol nhận group message."],
            ["6", "Broadcast", "Toàn bộ peer online nhận message."],
            ["7", "File transfer", "Bob nhận file README.md và metadata được lưu."],
            ["8", "Carol offline", "Message Alice -> Carol bị failed_queued."],
            ["9", "Carol online lại", "Carol nhận message delivered_from_queue."],
        ],
        widths=[1.2, 6.4, 8.0],
    )
    doc.add_heading("7.2. Kết quả direct, group, broadcast", level=2)
    add_figure(doc, shots["peer_b_received"], "Peer Bob nhận direct, group, broadcast và file event")
    add_figure(doc, shots["peer_c_delivered"], "Peer Carol nhận group và offline message sau khi online lại")
    add_para(doc, "Kết quả cho thấy Bob và Carol nhận message realtime. Với direct message, Bob nhận tin từ Alice và hệ thống trả ACK. Với group message, Alice gửi riêng tới Bob và Carol; mỗi peer xử lý payload và cập nhật message list của chính mình.")
    doc.add_heading("7.3. Kết quả file transfer", level=2)
    add_figure(doc, shots["peer_b_file"], "Peer Bob sau khi nhận file README.md từ Alice")
    add_para(doc, "File README.md được gửi bằng multipart từ Web/API của Alice, sau đó runtime đóng gói thành file_transfer payload và gửi tới Bob qua TCP. Bob nhận payload, trả ACK và ghi log file received. Vì file đang đóng gói base64, cơ chế này phù hợp file nhỏ và dễ demo.")
    doc.add_heading("7.4. Kết quả offline queue", level=2)
    add_figure(doc, shots["peer_a_offline"], "Alice gửi tin khi Carol offline và message bị queued")
    add_figure(doc, shots["peer_c_delivered"], "Carol online lại và nhận message từ offline queue")
    add_para(doc, "Tình huống offline queue chứng minh hệ thống có khả năng chịu lỗi cơ bản. Khi TCP connection tới Carol bị từ chối, Alice retry ba lần, sau đó lưu payload vào bootstrap offline queue. Khi Carol khởi động lại, runtime poll queue và nhận message với trạng thái delivered_from_queue.")
    doc.add_heading("7.5. Kiểm thử churn", level=2)
    add_figure(doc, diagrams["churn"], "Churn simulation mô phỏng peer rời mạng và tham gia lại")
    add_para(doc, "Churn simulation hỗ trợ kiểm tra thay đổi trạng thái online/offline theo thời gian. Khi churn chuyển peer sang offline, peer unregister khỏi bootstrap và dừng TCP server. Khi quay lại online, peer start TCP server, register lại và sync dữ liệu.")
    doc.add_heading("7.6. Đánh giá kết quả", level=2)
    add_bullets(
        doc,
        [
            "Hệ thống đáp ứng chức năng cốt lõi của đề tài chat P2P.",
            "Luồng chat chính sử dụng TCP peer-to-peer, không chuyển tiếp qua bootstrap.",
            "Web UI trực quan, thể hiện trạng thái delivery và log phù hợp cho bảo vệ.",
            "ACK, retry và offline queue hoạt động đúng trong tình huống peer offline.",
            "Database schema đầy đủ cho peer, group, message, ACK, file transfer và log.",
            "Project có thể chạy nhanh bằng memory fallback và có thể chuyển sang MySQL khi cần.",
        ],
    )


def chapter_8(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG 8. KẾT LUẬN, HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN", level=1)
    doc.add_heading("8.1. Kết luận", level=2)
    for t in [
        "Đồ án đã xây dựng được hệ thống chat ngang hàng P2P bằng Node.js, Express, TCP socket, Socket.IO và MySQL. Mỗi peer là một node độc lập, có Web UI, TCP server, TCP client và runtime điều phối riêng.",
        "Bootstrap server được sử dụng đúng vai trò hỗ trợ discovery, heartbeat, group metadata, offline queue và launcher. Luồng chat chính vẫn đi trực tiếp giữa các peer qua TCP socket. Đây là điểm quan trọng chứng minh hệ thống không phải chat client-server truyền thống.",
        "Hệ thống đã triển khai các chức năng cơ bản và nâng cao: direct chat, group chat, broadcast, ACK, retry, timeout, store-and-forward, file transfer, mã hóa AES-256-GCM, Web UI realtime, lưu history và mô phỏng churn.",
    ]:
        add_para(doc, t)
    doc.add_heading("8.2. Hạn chế", level=2)
    add_bullets(
        doc,
        [
            "Chưa hỗ trợ NAT traversal nên phù hợp chạy local hoặc LAN hơn là Internet production.",
            "Chưa có hệ thống đăng nhập, phân quyền và xác thực peer đầy đủ.",
            "Encryption dùng shared secret chung, chưa có key exchange riêng cho từng peer.",
            "File transfer đang phù hợp file nhỏ vì dữ liệu được đóng gói base64 trong payload.",
            "Bootstrap vẫn là điểm hỗ trợ quan trọng cho discovery và offline queue.",
        ],
    )
    doc.add_heading("8.3. Hướng phát triển", level=2)
    add_bullets(
        doc,
        [
            "Bổ sung JWT/session để xác thực người dùng và quản lý quyền truy cập.",
            "Dùng public/private key cho từng peer và triển khai key exchange an toàn.",
            "Tối ưu file transfer bằng streaming hoặc chia chunk, có hash kiểm tra toàn vẹn.",
            "Nghiên cứu WebRTC DataChannel hoặc NAT traversal để chạy qua Internet.",
            "Đóng gói bằng Docker Compose cho MySQL, bootstrap và nhiều peer.",
            "Bổ sung test tự động cho TCP protocol, REST API và luồng offline queue.",
            "Thêm dashboard thống kê latency, tỷ lệ delivery, retry và queue.",
        ],
    )


def references_and_appendices(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        "Node.js Documentation: https://nodejs.org/docs",
        "Express Documentation: https://expressjs.com/",
        "Socket.IO Documentation: https://socket.io/docs/",
        "MySQL Documentation: https://dev.mysql.com/doc/",
        "OWASP Cryptographic Storage Cheat Sheet: https://cheatsheetseries.owasp.org/",
        "Tanenbaum, A. S. và Van Steen, M. Distributed Systems: Principles and Paradigms.",
    ]
    add_numbered(doc, refs)

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC A. HƯỚNG DẪN CHẠY HỆ THỐNG", level=1)
    add_code_block(
        doc,
        """
npm install
npm run bootstrap
npm run peer:a
npm run peer:b
npm run peer:c
npm run check
""",
        "Lệnh chạy nhanh",
    )
    add_numbered(
        doc,
        [
            "Mở http://127.0.0.1:3000 để xem Bootstrap Launcher.",
            "Mở http://127.0.0.1:3101 để xem Alice.",
            "Mở http://127.0.0.1:3102 để xem Bob.",
            "Mở http://127.0.0.1:3103 để xem Carol.",
            "Từ Alice gửi direct message tới Bob.",
            "Tạo group gồm Bob và Carol rồi gửi group message.",
            "Gửi broadcast tới toàn mạng.",
            "Stop Carol, gửi message tới Carol, sau đó start Carol lại để kiểm tra offline queue.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC B. API VÀ CẤU HÌNH", level=1)
    add_table(
        doc,
        "Danh sách API Bootstrap",
        ["Endpoint", "Chức năng"],
        [
            ["GET /health", "Kiểm tra bootstrap server."],
            ["POST /api/register", "Đăng ký peer."],
            ["POST /api/unregister", "Đánh dấu peer offline."],
            ["POST /api/heartbeat", "Cập nhật last_seen."],
            ["GET /api/peers", "Lấy danh sách peer."],
            ["POST /api/groups", "Tạo group."],
            ["GET /api/offline-messages/:peerId", "Lấy offline messages."],
            ["POST /api/acks", "Lưu ACK."],
            ["POST /api/file-transfers", "Lưu metadata file transfer."],
        ],
        widths=[6, 9.5],
    )
    add_table(
        doc,
        "Danh sách API Peer",
        ["Endpoint", "Chức năng"],
        [
            ["GET /api/me", "Lấy thông tin peer hiện tại."],
            ["GET /api/peers", "Lấy peer cache."],
            ["POST /api/messages/direct", "Gửi direct message."],
            ["POST /api/messages/group", "Gửi group message."],
            ["POST /api/broadcast", "Gửi broadcast."],
            ["POST /api/groups", "Tạo group từ peer UI."],
            ["POST /api/files", "Gửi file."],
            ["POST /api/churn/start", "Bật churn simulation."],
            ["POST /api/sync", "Sync peers, groups và offline queue."],
        ],
        widths=[6, 9.5],
    )

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC C. KỊCH BẢN BẢO VỆ", level=1)
    add_para(doc, "Khi trình bày trước giảng viên, nên nhấn mạnh ba ý: hệ thống đúng bản chất P2P, bootstrap chỉ hỗ trợ discovery/metadata, và hệ thống có xử lý lỗi cơ bản bằng ACK, retry, timeout và offline queue.")
    add_table(
        doc,
        "Câu hỏi bảo vệ thường gặp",
        ["Câu hỏi", "Gợi ý trả lời"],
        [
            ["Hệ thống P2P ở đâu?", "Mỗi peer mở TCP server riêng và message chính đi trực tiếp giữa các peer qua TCP."],
            ["Bootstrap có phải server trung tâm không?", "Bootstrap là server hỗ trợ discovery và metadata, không chuyển tiếp luồng chat chính."],
            ["Nếu peer nhận bị tắt thì sao?", "Peer gửi retry, nếu vẫn lỗi thì queue offline; peer nhận online lại sẽ poll queue."],
            ["ACK dùng để làm gì?", "ACK xác nhận peer nhận đã xử lý message để peer gửi cập nhật delivered."],
            ["Chat nhóm hoạt động thế nào?", "Peer gửi lấy danh sách member rồi gửi TCP riêng tới từng thành viên."],
            ["Vì sao vẫn có MySQL?", "MySQL lưu trạng thái và lịch sử, không phải kênh truyền message realtime."],
        ],
        widths=[5.2, 10.3],
    )
    add_para(doc, "Câu chốt đề xuất: Project đã mô phỏng được hệ thống chat ngang hàng P2P với nhiều peer độc lập, giao tiếp trực tiếp bằng TCP, có discovery, ACK, retry, mã hóa, group, broadcast, file transfer và offline queue.")


PREDECLARED_FIGURES = [
    "Kiến trúc tổng quan giữa bootstrap và các peer",
    "Sơ đồ triển khai local với Bootstrap và ba peer demo",
    "Các module chính trong Bootstrap Server",
    "Các module chính trong Peer App",
    "Sequence direct message, ACK và lưu lịch sử",
    "Thiết kế group message gửi riêng tới từng thành viên",
    "Thiết kế broadcast tới toàn bộ peer online",
    "Store-and-forward khi peer đích offline",
    "Máy trạng thái gửi tin và xử lý ACK",
    "ERD các bảng chính của hệ thống",
    "Luồng mã hóa và giải mã AES-256-GCM",
    "Luồng gửi file qua TCP payload",
    "Cơ chế MySQL và memory fallback",
    "Bootstrap Launcher hiển thị peer đang online",
    "Peer Alice sau khi gửi direct, group và broadcast",
    "Tab Group trên Peer Alice",
    "Tab Broadcast trên Peer Alice",
    "Tab File transfer trên Peer Alice",
    "Bản đồ API Bootstrap và Peer",
    "Ma trận kiểm thử chức năng chính",
    "Peer Bob nhận direct, group, broadcast và file event",
    "Peer Carol nhận group và offline message sau khi online lại",
    "Peer Bob sau khi nhận file README.md từ Alice",
    "Alice gửi tin khi Carol offline và message bị queued",
    "Carol online lại và nhận message từ offline queue",
    "Churn simulation mô phỏng peer rời mạng và tham gia lại",
]

PREDECLARED_TABLES = [
    "Thông tin nhiệm vụ và phạm vi báo cáo",
    "Từ khóa chính của đề tài",
    "Danh mục từ viết tắt",
    "Phạm vi chức năng của đề tài",
    "Stack công nghệ sử dụng trong project",
    "Yêu cầu chức năng chi tiết",
    "Tác nhân trong hệ thống",
    "Tiêu chí nghiệm thu đồ án",
    "Mô tả các bảng dữ liệu chính",
    "Các loại payload chính",
    "Scripts npm phục vụ chạy và kiểm tra",
    "Kịch bản kiểm thử chức năng",
    "Danh sách API Bootstrap",
    "Danh sách API Peer",
    "Câu hỏi bảo vệ thường gặp",
]


def build() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    diagrams = make_diagrams()
    shots = prepare_screenshots()

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    chapter_1(doc)
    chapter_2(doc, diagrams)
    chapter_3(doc)
    chapter_4(doc, diagrams)
    chapter_5(doc, diagrams)
    chapter_6(doc, diagrams, shots)
    chapter_7(doc, diagrams, shots)
    chapter_8(doc)
    references_and_appendices(doc)

    core = doc.core_properties
    core.title = "Báo cáo đồ án tốt nghiệp - Hệ thống chat ngang hàng P2P"
    core.subject = "P2P Chat System"
    core.keywords = "P2P, chat, TCP, Socket.IO, MySQL, distributed systems"
    core.comments = "Generated from project source and demo screenshots."
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    output = build()
    print(str(output).encode("unicode_escape").decode("ascii"))
